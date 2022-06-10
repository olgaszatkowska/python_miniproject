import shutil
import numpy
import requests
import pandas
import seaborn as sns
import matplotlib.pyplot as plt
from PIL import Image
import os
import math
from docx import Document
from docx.shared import Inches

from src.chapter_info import ChapterData
from src.book_info import BookInfo


def _find(something):
    def _(line):
        return something in line

    return _


def _parse_value(line):
    return line[0].split(":")[1].strip()


def _find_index_of(book, word):
    for line in book:
        if word.lower() in line.lower():
            yield book.index(line)


def _find_fist_chapter(book):
    index_generator = _find_index_of(book, "chapter")
    first_chapter = next(index_generator)
    second_chapter = next(index_generator)
    return book[first_chapter + 2 : second_chapter]


def create_book_info(book):
    author = list(filter(_find("Author:"), book))
    title = list(filter(_find("Title:"), book))
    first_chapter = _find_fist_chapter(book)
    return BookInfo(_parse_value(title), _parse_value(author), first_chapter)


def parse_book_as_list(book_url):
    response = requests.get(book_url)
    return [line.strip() for line in response.text.splitlines()]


def _roundup(x):
    return int(math.ceil(x / 10.0)) * 10


def _count_words(paragraph):
    raw_len = len(paragraph.split())
    return _roundup(raw_len)


def get_words_per_paragraph(chapter):
    paragraphs = []
    paragraph = ""
    for line in chapter:
        paragraph = f"{paragraph}{line}" if "*" not in line else ""
        if line == "" and paragraph != "":
            paragraphs.append(_count_words(paragraph))
            paragraph = ""
    return pandas.DataFrame(
        data=[[v + 1, k] for v, k in enumerate(paragraphs)],
        columns=["Paragraph", "Words"],
    )


def generate_plot(df, filename):
    x, y = df.columns
    ax = sns.barplot(x=x, y=y, palette="ch:.25", data=df)
    ax.set(title="Words per paragraph")
    ax.set_xticklabels(ax.get_xticklabels(), fontsize=5)
    plt.savefig(filename)


def download_picture(url):
    filename = url.split("/")[-1]
    response = requests.get(url, stream=True)
    if response.status_code == 200:
        response.raw.decode_content = True
        with open(filename, "wb") as f:
            shutil.copyfileobj(response.raw, f)
    return filename


def save_image(image, path, beginning=""):
    *filename, extension = path.split(".")
    joined_filename = "".join(filename)
    new_path = f"{beginning}{joined_filename}.{extension}"
    image.save(new_path)
    return new_path


def crop_image(path, size):
    cropped = Image.open(path).crop(size)
    return save_image(cropped, path, "cropped_")


def rotate_image(path, degree):
    rotated = Image.open(path).rotate(degree)
    return save_image(rotated, path, "rotated_")


def compose_images(foreground_path, background_path, place=(0, 0)):
    foreground_img = Image.open(foreground_path)
    background_img = Image.open(background_path)
    background_img.paste(foreground_img, place, foreground_img)
    save_image(background_img, background_path, "composed_")


def create_word_document(
    filename, author, book_title, cover_path, book_info, graph_path, chapter_data
):
    document = Document()
    document.add_heading(book_title, 0)
    document.add_picture(cover_path, width=Inches(5))
    p_author = document.add_paragraph("Author of the book: ")
    p_author.add_run(book_info.author).bold = True
    p_author.add_run(". Author of document: ")
    p_author.add_run(author).italic = True
    document.add_picture(graph_path, width=Inches(5))
    document.add_paragraph(str(chapter_data))
    document.save(filename)


def create_chapter_data(chapter_df):
    word_count = chapter_df.values[:, 1]
    return ChapterData(
        len(word_count),
        sum(word_count),
        word_count.min(),
        word_count.max(),
        round(numpy.average(word_count)),
    )


def remove_files(*args):
    for file in args:
        os.remove(file)
